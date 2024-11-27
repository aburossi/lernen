import streamlit as st
import time
from openai import OpenAI
import json
from PyPDF2 import PdfReader
from fpdf import FPDF
from io import BytesIO
from docx import Document
from streamlit_cookies_manager import EncryptedCookieManager

# Set page config
st.set_page_config(page_title="Exam Creator", page_icon="📝")

__version__ = "1.2.0"

def stream_llm_response(messages, model_params, api_key):
    client = OpenAI(api_key=api_key)
    response = client.chat.completions.create(
        model=model_params["model"],  # Use the passed model
        messages=messages,
        temperature=model_params.get("temperature", 0.5),
        max_tokens=15096,
    )
    return response.choices[0].message.content

def extract_text_from_pdf(pdf_file):
    pdf_reader = PdfReader(pdf_file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text() + "\n"
    return text

def chunk_text(text, max_tokens=3000):
    sentences = text.split('. ')
    chunks = []
    chunk = ""
    for sentence in sentences:
        if len(chunk) + len(sentence) > max_tokens:
            chunks.append(chunk)
            chunk = sentence + ". "
        else:
            chunk += sentence + ". "
    if chunk:
        chunks.append(chunk)
    return chunks

def generate_mc_questions(content_text, api_key, model_key):
    system_prompt = (
        "Sie sind ein Lehrer für Allgemeinbildung und sollen eine Prüfung zum Thema des eingereichten PDFs erstellen. "
        "Verwenden Sie den Inhalt des PDFs (bitte gründlich analysieren) und erstellen Sie eine Multiple-Choice-Prüfung auf Oberstufenniveau. "
        "Die Prüfung soll sowohl Fragen mit einer richtigen Antwort als auch Fragen mit mehreren richtigen Antworten enthalten. "
        "Kennzeichnen Sie die Fragen entsprechend, damit die Schüler wissen, wie viele Optionen sie auswählen sollen. "
        "Erstellen Sie so viele Prüfungsfragen, wie nötig sind, um den gesamten Inhalt abzudecken, aber maximal 20 Fragen. "
        "Geben Sie die Ausgabe im JSON-Format an. "
        "Das JSON sollte folgende Struktur haben: [{'question': '...', 'choices': ['...'], 'correct_answer': '...', 'explanation': '...'}, ...]. "
        "Stellen Sie sicher, dass das JSON gültig und korrekt formatiert ist."
    )
    user_prompt = (
        "Using the following content from the uploaded PDF, create multiple-choice and single-choice questions. "
        "Ensure that each question is based on the information provided in the PDF content. "
        "Mark questions appropriately so that students know how many options to select. "
        "Create as many questions as necessary to cover the entire content, but no more than 20 questions. "
        "Provide the output in JSON format with the following structure: "
        "[{'question': '...', 'choices': ['...'], 'correct_answer': '...', 'explanation': '...'}, ...]. "
        "Ensure the JSON is valid and properly formatted.\n\nPDF Content:\n\n"
    ) + content_text

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt},
    ]
    try:
        response = stream_llm_response(messages, model_params={"model": model_key, "temperature": 0.5}, api_key=api_key)
        return response, None
    except Exception as e:
        return None, str(e)


def parse_generated_questions(response):
    try:
        json_start = response.find('[')
        json_end = response.rfind(']') + 1
        if json_start == -1 or json_end == 0:
            return None, f"No JSON data found in the response. First 500 characters of response:\n{response[:500]}..."
        json_str = response[json_start:json_end]

        questions = json.loads(json_str)
        return questions, None
    except json.JSONDecodeError as e:
        return None, f"JSON parsing error: {e}\n\nFirst 500 characters of response:\n{response[:500]}..."
    except Exception as e:
        return None, f"Unexpected error: {str(e)}\n\nFirst 500 characters of response:\n{response[:500]}..."

class PDF(FPDF):
    def header(self):
        self.set_font('Arial', 'B', 12)
        self.cell(0, 10, 'Generated Exam', 0, 1, 'C')

    def chapter_title(self, title):
        self.set_font('Arial', 'B', 12)
        self.multi_cell(0, 10, title)
        self.ln(5)

    def chapter_body(self, body):
        self.set_font('Arial', '', 12)
        self.multi_cell(0, 10, body)
        self.ln()

def generate_pdf(questions):
    pdf = PDF()

    # Part 1: Questions without solutions
    pdf.add_page()
    pdf.set_font("Arial", "B", 16)
    pdf.cell(0, 10, "Exam - Questions Only (No Solutions)", 0, 1, "C")
    pdf.ln(10)

    for i, q in enumerate(questions):
        question = f"Q{i+1}: {q['question']}"
        pdf.chapter_title(question)

        # Correctly format the choices
        choices = "\n".join([f"{chr(65+j)}. {choice}" for j, choice in enumerate(q['choices'])])
        pdf.chapter_body(choices)

    # Part 2: Questions with solutions
    pdf.add_page()
    pdf.set_font("Arial", "B", 16)
    pdf.cell(0, 10, "Exam - Questions with Solutions", 0, 1, "C")
    pdf.ln(10)

    for i, q in enumerate(questions):
        question = f"Q{i+1}: {q['question']}"
        pdf.chapter_title(question)

        # Correctly format the choices
        choices = "\n".join([f"{chr(65+j)}. {choice}" for j, choice in enumerate(q['choices'])])
        pdf.chapter_body(choices)

        correct_answer = f"Correct answer: {q['correct_answer']}"
        pdf.chapter_body(correct_answer)

        explanation = f"Explanation: {q['explanation']}"
        pdf.chapter_body(explanation)

    return pdf.output(dest="S").encode("latin1")

def generate_docx(questions):
    doc = Document()

    # Part 1: Questions without solutions
    doc.add_heading('Exam - Questions Only (No Solutions)', level=1)

    for i, q in enumerate(questions):
        question = f"Q{i+1}: {q['question']}"
        doc.add_heading(question, level=2)

        # Add choices
        for j, choice in enumerate(q['choices']):
            doc.add_paragraph(f"{chr(65+j)}. {choice}")

    # Add a page break
    doc.add_page_break()

    # Part 2: Questions with solutions
    doc.add_heading('Exam - Questions with Solutions', level=1)

    for i, q in enumerate(questions):
        question = f"Q{i+1}: {q['question']}"
        doc.add_heading(question, level=2)

        # Add choices
        for j, choice in enumerate(q['choices']):
            doc.add_paragraph(f"{chr(65+j)}. {choice}")

        # Add correct answer and explanation
        doc.add_paragraph(f"Correct answer: {q['correct_answer']}")
        doc.add_paragraph(f"Explanation: {q['explanation']}")

    # Save the document to a BytesIO object
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)  # Reset stream position for reading
    return file_stream.getvalue()

def main():
    # Set up the sidebar
    st.sidebar.title("Exam Creator Sidebar")
    st.sidebar.write("Navigate through the app and access additional resources here.")

    # Dropdown for app mode
    app_mode_options = ["Upload PDF & Generate Questions", "Take the Quiz", "Download as PDF"]
    if "app_mode" not in st.session_state:
        st.session_state.app_mode = app_mode_options[0]  # Default to the first option
    st.session_state.app_mode = st.sidebar.selectbox("Choose the app mode", app_mode_options)

    # Embed the video link
    st.sidebar.markdown("[How to get your OpenAI API Key](https://youtu.be/NsTAjBdHb1k)")

    # License and contact information
    st.sidebar.subheader("License & Contact")
    st.sidebar.write("This application is licensed for personal and educational use.")
    st.sidebar.markdown("For inquiries, email: [pietro.rossi@bbw.ch](mailto:pietro.rossi@bbw.ch)")

    # Initialize the cookie manager
    cookies = EncryptedCookieManager(prefix="exam_creator_app")  # Prefix to avoid conflicts
    if not cookies.ready():
        st.stop()
    
    # API Key input and cookie management
    if "api_key" not in st.session_state:
        st.session_state.api_key = cookies.get("api_key", "")
    
    api_key_input = st.text_input(
        "Enter your OpenAI API Key:",
        value=st.session_state.api_key,
        type="password",
        help="Your API key will be saved in your browser securely if you choose to remember it."
    )
    
    # Save API Key in session state
    if api_key_input:
        st.session_state.api_key = api_key_input
    
    # Save or Clear API Key
    col1, col2 = st.sidebar.columns(2)
    with col1:
        if st.button("Remember API Key"):
            cookies["api_key"] = st.session_state.api_key
            cookies.save()  # Save the cookie
            st.success("API Key saved to your browser.")
    with col2:
        if st.button("Forget API Key"):
            if "api_key" in cookies:
                del cookies["api_key"]
                cookies.save()  # Save changes to delete the cookie
                st.session_state.api_key = ""
                st.success("API Key removed from your browser.")
    
    # Use the API key in the app
    api_key = st.session_state.api_key
    if not api_key:
        st.warning("Please provide your OpenAI API key to proceed.")


    # Model selection
    model_options = {
        "gpt-4o-mini (cheapest, fastest)": "gpt-4o-mini",
        "gpt-4o (better output)": "gpt-4o"
    }
    selected_model = st.sidebar.radio(
        "Select GPT model",
        options=list(model_options.keys()),
        index=0  # Default to "gpt-4o-mini (cheapest, fastest)"
    )
    selected_model_key = model_options[selected_model]

    # Load the appropriate app mode
    if st.session_state.app_mode == "Upload PDF & Generate Questions":
        pdf_upload_app(api_key, selected_model_key)
    elif st.session_state.app_mode == "Take the Quiz":
        if 'mc_test_generated' in st.session_state and st.session_state.mc_test_generated:
            if 'generated_questions' in st.session_state and st.session_state.generated_questions:
                mc_quiz_app()
            else:
                st.warning("No generated questions found. Please upload a PDF and generate questions first.")
        else:
            st.warning("Please upload a PDF and generate questions first.")
    elif st.session_state.app_mode == "Download as PDF":
        download_pdf_app()

def pdf_upload_app(api_key, selected_model_key):
    st.subheader("Upload Your Content - Create Your Test Exam")
    st.write("Upload the content and we take care of the rest")

    content_text = ""

    if 'messages' not in st.session_state:
        st.session_state.messages = []

    uploaded_pdf = st.file_uploader("Upload a PDF document", type=["pdf"])
    if uploaded_pdf and api_key:
        pdf_text = extract_text_from_pdf(uploaded_pdf)
        content_text += pdf_text
        st.success("PDF content added to the session.")

        # Display a sample of the extracted text for verification
        st.subheader("Sample of extracted PDF content:")
        st.text(content_text[:500] + "...")  # Display first 500 characters

        st.info("Generating the exam from the uploaded content. It will take just a minute...")
        chunks = chunk_text(content_text)
        questions = []
        for chunk in chunks:
            response, error = generate_mc_questions(chunk, api_key, selected_model_key)
            if error:
                st.error(f"Error generating questions: {error}")
                break
            parsed_questions, parse_error = parse_generated_questions(response)
            if parse_error:
                st.error(parse_error)
                st.text("Full response:")
                st.text(response)
                break
            if parsed_questions:
                questions.extend(parsed_questions)
                if len(questions) >= 20:
                    questions = questions[:20]  # Limit to 20 questions
                    break
        if questions:
            st.session_state.generated_questions = questions
            st.session_state.content_text = content_text
            st.session_state.mc_test_generated = True
            st.success(f"The exam has been successfully created with {len(questions)} questions! Switch the Sidebar Panel to take the exam.")

            # Display a sample question for verification
            st.subheader("Sample generated question:")
            st.json(questions[0])

            time.sleep(2)
            st.session_state.app_mode = "Take the Quiz"
            st.rerun()
        else:
            st.error("No questions were generated. Please check the error messages above and try again.")
    elif not api_key:
        st.warning("Please enter your OpenAI API key.")
    else:
        st.warning("Please upload a PDF to generate the interactive exam.")

def submit_answer(i, quiz_data):
    user_choice = st.session_state[f"user_choice_{i}"]
    st.session_state.answers[i] = user_choice
    if user_choice == quiz_data['correct_answer']:
        st.session_state.feedback[i] = ("Correct", quiz_data.get('explanation', 'No explanation available'))
        st.session_state.correct_answers += 1
    else:
        st.session_state.feedback[i] = ("Incorrect", quiz_data.get('explanation', 'No explanation available'), quiz_data['correct_answer'])

def mc_quiz_app():
    st.subheader("Multiple Choice Exam")
    st.write("There is always one correct answer per question.")

    # Buttons for downloading PDF and DOCX
    col1, col2 = st.columns(2)

    with col1:
        if st.button("Download the exam as PDF"):
            if 'generated_questions' in st.session_state:
                pdf_bytes = generate_pdf(st.session_state.generated_questions)
                st.download_button(
                    label="Download PDF",
                    data=pdf_bytes,
                    file_name="generated_exam.pdf",
                    mime="application/pdf"
                )
            else:
                st.warning("No questions are available for download.")

    with col2:
        if st.button("Download the exam as DOCX"):
            if 'generated_questions' in st.session_state:
                docx_bytes = generate_docx(st.session_state.generated_questions)
                st.download_button(
                    label="Download DOCX",
                    data=docx_bytes,
                    file_name="generated_exam.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                st.warning("No questions are available for download.")

    questions = st.session_state.generated_questions

    if questions:
        if 'answers' not in st.session_state:
            st.session_state.answers = [None] * len(questions)
            st.session_state.feedback = [None] * len(questions)
            st.session_state.correct_answers = 0

        for i, quiz_data in enumerate(questions):
            st.markdown(f"### Question {i+1}: {quiz_data['question']}")

            if st.session_state.answers[i] is None:
                user_choice = st.radio("Choose an answer:", quiz_data['choices'], key=f"user_choice_{i}")
                st.button(f"Submit your answer {i+1}", key=f"submit_{i}", on_click=submit_answer, args=(i, quiz_data))
            else:
                st.radio("Choose an answer:", quiz_data['choices'], key=f"user_choice_{i}", index=quiz_data['choices'].index(st.session_state.answers[i]), disabled=True)
                if st.session_state.feedback[i][0] == "Correct":
                    st.success(st.session_state.feedback[i][0])
                else:
                    st.error(f"{st.session_state.feedback[i][0]} - Correct answer: {st.session_state.feedback[i][2]}")
                st.markdown(f"Explanation: {st.session_state.feedback[i][1]}")

        if all(answer is not None for answer in st.session_state.answers):
            score = st.session_state.correct_answers
            total_questions = len(questions)
            st.write(f"""
                <div style="display: flex; flex-direction: column; align-items: center; justify-content: center; height: 100vh;">
                    <h1 style="font-size: 3em; color: gold;">🏆</h1>
                    <h1>Your Score: {score}/{total_questions}</h1>
                </div>
            """, unsafe_allow_html=True)


def download_pdf_app():
    st.subheader('Download Your Exam as PDF')

    questions = st.session_state.generated_questions

    if questions:
        for i, q in enumerate(questions):
            st.markdown(f"### Q{i+1}: {q['question']}")
            for choice in q['choices']:
                st.write(choice)
            st.write(f"**Correct answer:** {q['correct_answer']}")
            st.write(f"**Explanation:** {q['explanation']}")
            st.write("---")

        pdf_bytes = generate_pdf(questions)
        st.download_button(
            label="Download PDF",
            data=pdf_bytes,
            file_name="generated_exam.pdf",
            mime="application/pdf"
        )

if __name__ == '__main__':
    main()


